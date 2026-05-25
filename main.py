VERSION_LABEL = "v9_professional_reminder_notes"

import os
import json
import random
from datetime import datetime, timedelta
from calendar import monthrange
from zoneinfo import ZoneInfo

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, BotCommand, InputFile
from telegram.ext import Application, CommandHandler, CallbackQueryHandler, ContextTypes, MessageHandler, filters

try:
    from docx import Document
except Exception:
    Document = None

try:
    from openpyxl import Workbook
except Exception:
    Workbook = None

TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
TZ = ZoneInfo("America/Vancouver")

CONFIG_FILE = "life_config.json"
DATA_FILE = "life_data.json"
NOTES_FILE = "notes.json"
REMINDERS_FILE = "reminders.json"
USER_STATE_FILE = "user_state.json"
EXPORT_DIR = "exports"

DEFAULT_CONFIG = {
    "study": {"title": "📚 مطالعه", "subjects": {
        "cpa": {"title": "CPA", "subs": {"fr": "Financial Reporting / مالی", "tax": "Taxation / مالیات", "audit": "Audit / حسابرسی", "mgmt": "Management / مدیریت"}},
        "english": {"title": "English / انگلیسی", "subs": {"grammar": "Grammar", "speaking": "Speaking", "listening": "Listening", "vocab": "Vocabulary"}},
        "python": {"title": "Python / پایتون", "subs": {"basic": "Basics", "finance": "Finance Analysis", "bot": "Telegram Bot"}}
    }},
    "trading": {"title": "📊 ترید عملی", "subjects": {
        "xauusd": {"title": "XAUUSD / طلا", "subs": {"analysis": "تحلیل", "live": "معامله زنده", "journal": "ژورنال معامله", "backtest": "بک‌تست"}},
        "eurusd": {"title": "EURUSD", "subs": {"analysis": "تحلیل", "live": "معامله زنده", "journal": "ژورنال معامله", "backtest": "بک‌تست"}},
        "usdjpy": {"title": "USDJPY", "subs": {"analysis": "تحلیل", "live": "معامله زنده", "journal": "ژورنال معامله", "backtest": "بک‌تست"}},
        "nasdaq": {"title": "NASDAQ", "subs": {"analysis": "تحلیل", "live": "معامله زنده", "journal": "ژورنال معامله", "backtest": "بک‌تست"}},
        "btc": {"title": "BTC / کریپتو", "subs": {"analysis": "تحلیل", "live": "معامله زنده", "journal": "ژورنال معامله", "backtest": "بک‌تست"}}
    }},
    "uber": {"title": "🚗 اوبر", "subjects": {"drive": {"title": "Driving / رانندگی", "subs": {"work": "Work Shift", "airport": "Airport", "city": "City"}}}},
    "walk": {"title": "🚶 پیاده‌روی", "subjects": {"walk": {"title": "Walking", "subs": {"normal": "Normal Walk", "fast": "Fast Walk", "family": "Family Walk"}}}},
    "sleep": {"title": "😴 خواب", "subjects": {"sleep": {"title": "Sleep", "subs": {"night": "Night Sleep", "nap": "Nap"}}}},
    "misc": {"title": "🧩 متفرقه", "subjects": {
        "shopping": {"title": "🛒 خرید", "subs": {"grocery": "خرید خانه", "personal": "خرید شخصی", "car": "خرید ماشین/ابزار"}},
        "errands": {"title": "📌 کارهای بیرون", "subs": {"bank": "بانک", "school": "مدرسه", "appointment": "قرار/ملاقات"}},
        "home": {"title": "🏠 خانه", "subs": {"clean": "نظافت", "repair": "تعمیرات", "family": "خانواده"}}
    }}
}

MIND_MESSAGES = {
    "mind": [
        "🧠 فریب ذهن\n\nذهن می‌خواهد دائم در بازار باشد؛ اما بازار کار خودش را می‌کند. تو نباید دائم در بازار باشی.",
        "🧠 شمع بلند\n\nذهن به‌راحتی فریب یک شمع بلند را می‌خورد. شمع بلند دلیل ورود نیست.",
        "🧠 سیستم یا احساس؟\n\nهر وقت حس کردی باید همین الآن معامله کنی، اول بپرس: این سیستم است یا اضطراب؟",
        "🧠 چارت\n\nچارت را باز کن، اما اجازه نده چارت تو را باز کند.",
        "🧠 خاطره دیروز\n\nترس امروز، گاهی خاطره دیروز است، نه حقیقت بازار امروز.",
        "🧠 ندانستن\n\nذهن می‌خواهد زود تصمیم بگیرد، چون از ندانستن می‌ترسد. تریدر حرفه‌ای ندانستن را تحمل می‌کند."
    ],
    "business": [
        "🏪 والمارت خانگی\n\nتو یک والمارت خانگی داری. والمارت هم همه موجودی‌اش را در یک ساعت نمی‌فروشد.",
        "🏪 مغازه حساب\n\nحساب معاملاتی تو مغازه توست. در یک ساعت مغازه را آتش نزن. فردا هم مشتری هست.",
        "💼 بیزینس ترید\n\nترید یک بیزینس است. هر معامله یک تصمیم مدیریتی است، نه هیجان لحظه‌ای.",
        "💼 گردش بزرگ\n\nروزی یک تا ده میلیون دلار می‌توانی معامله کنی؛ اما نه در یک نقطه، نه در یک قیمت، نه در یک لحظه.",
        "🛡 حفظ مغازه\n\nهدف اول فروش بیشتر نیست؛ هدف اول باز ماندن مغازه است.",
        "🏪 والمارت\n\nوالمارت همه اجناسش را یک‌جا نمی‌فروشد. تو هم لازم نیست همه فرصت‌های بازار را امروز بگیری."
    ],
    "treasure": [
        "🏴‍☠️ گنج بازار\n\nبازار یک گنج عظیم است؛ اما فقط برای کسی که نقشه دارد.",
        "🏴‍☠️ برداشت از گنج\n\nبرای گنج وارد می‌شوم؛ نه برای جنگ. سهمم را برمی‌دارم و برمی‌گردم.",
        "🏴‍☠️ کشتی\n\nحساب تو کشتی توست. برای یک سکه، کشتی را غرق نکن.",
        "🏴‍☠️ نقشه گنج\n\nسیستم تو نقشه گنج توست. وقتی از نقشه خارج می‌شوی، وارد قلمرو ترس و طمع می‌شوی.",
        "🌊 اقیانوس\n\nبازار مثل اقیانوس است. کسی با دست و پا زدن، اقیانوس را شکست نداده.",
        "🏴‍☠️ جزیره گنج\n\nباید با احترام وارد شوی، برداری و برگردی. نباید در جزیره گنج چادر بزنی."
    ],
    "setup": [
        "📊 ستاپ کانال و میانگین\n\n۱. جهت تایم بالا مشخص شود.\n۲. کانال رسم شود.\n۳. میانگین‌ها بررسی شود.\n۴. ورود فقط روی اصلاح.\n۵. حد زیان کوچک.\n۶. خروج سریع.",
        "📐 کانال\n\nکف کانال در روند صعودی محل جستجوی خرید است؛ سقف کانال در روند نزولی محل جستجوی فروش.",
        "〽️ میانگین‌ها\n\nمیانگین‌ها زبان بازارند. واکنش قیمت را ببین؛ حدس نزن.",
        "📈 خرید حرفه‌ای\n\nدر روند صعودی، روی اصلاح و کف کانال خرید کن؛ نه وسط هیجان شمع بلند.",
        "📉 فروش حرفه‌ای\n\nدر روند نزولی، روی اصلاح و سقف کانال فروش کن؛ نه بعد از ریزش شدید.",
        "⚡ ورود درست\n\nورود خوب باید زود خودش را نشان دهد. اگر حرکت نکرد، شاید جای ورود اشتباه بوده."
    ],
    "risk": [
        "⚖️ حجم\n\nنه در یک نقطه، نه در یک قیمت، نه در یک لحظه، همه حجم را وارد نکن.",
        "🛑 حد زیان\n\nحد زیان یعنی: من اشتباه بودن را می‌پذیرم و حسابم را حفظ می‌کنم.",
        "📉 جبران\n\nبرای جبران یک اشتباه، اشتباه بزرگ‌تر نکن.",
        "⚖️ افزایش حجم\n\nهیچ‌وقت به معامله زیان‌ده برای جبران، حجم اضافه نکن.",
        "📉 زیان کوچک\n\nزیان را کوچک ببند. حساب زنده بماند، فرصت دوباره هست.",
        "📈 سود\n\nسود خوب کردی؟ چارت را ببند. برگرد به زندگی."
    ],
    "samsara": [
        "🔁 سامسارا\n\nضرر → ترس → جبران → حجم بیشتر → ضرر بیشتر.\nچرخه را بازار نساخته؛ ذهن ساخته است.",
        "🔁 توقف\n\nرهایی از چرخه یعنی بعد از ضرر، توقف؛ نه انتقام.",
        "🔁 ذهن قدیمی\n\nهر بار با همان ذهن قبلی وارد شوی، همان اشتباه قبلی دوباره متولد می‌شود.",
        "🕊 رهایی\n\nدیدن، پذیرفتن، عمل، رها کردن.",
        "🕊 مشاهده\n\nپنج دقیقه فقط نگاه کن. نه صعودی بگو، نه نزولی. فقط ببین."
    ]
}


EXTRA_MIND_MESSAGES = {
    "mind": [
        "🧠 ذهن معامله‌گر\n\nبازار در یک ساعت حساب را نابود نمی‌کند؛ ذهن بی‌قرار این کار را می‌کند.",
        "🧠 سیستم یا احساس؟\n\nورود باید از سیستم بیاید، نه از عجله، ترس، انتقام یا فریب یک شمع بلند.",
        "🧠 مشاهده\n\nپنج دقیقه فقط نگاه کن. نه صعودی بگو، نه نزولی. فقط ببین.",
        "🧠 ذهن داغ\n\nاگر بعد از سود یا ضرر هنوز ذهنت داغ است، معامله ممنوع.",
        "🧠 پذیرش\n\nبازار یا صعودی است، یا نزولی، یا رنج. وظیفه تو جنگیدن نیست؛ فقط دیدن است.",
        "🧠 ندانستن\n\nوقتی نمی‌دانی، ندانستن را بپذیر. عجله برای دانستن، آغاز خطاست."
    ],
    "business": [
        "🏪 والمارت خانگی\n\nتو یک والمارت خانگی داری. والمارت همه موجودی‌اش را در یک ساعت نمی‌فروشد.",
        "🏪 مغازه حساب\n\nحساب معاملاتی تو مغازه توست. در یک ساعت مغازه را آتش نزن. فردا هم مشتری هست.",
        "💼 گردش بزرگ\n\nروزی یک تا ده میلیون دلار می‌توانی معامله کنی؛ اما نه در یک نقطه، نه در یک قیمت، نه در یک لحظه.",
        "💼 بیزینس ترید\n\nهر معامله یک تصمیم مدیریتی است. مدیر خوب اول بقا را حفظ می‌کند، بعد سود می‌سازد.",
        "🛡 سرمایه\n\nسرمایه فقط عدد حساب نیست؛ اعتماد، بقا، تمرکز و ادامه کسب‌وکار توست.",
        "🎯 هدف\n\nهدف فقط سود نیست. هدف، داشتن و نگهداری از حساب است."
    ],
    "treasure": [
        "🏴‍☠️ گنج بازار\n\nبازار یک گنج عظیم است؛ اما فقط برای کسی که نقشه دارد.",
        "🏴‍☠️ قانون گنج\n\nبرای گنج وارد شو، نه برای جنگ. سهمت را بردار و برگرد.",
        "🏴‍☠️ کشتی\n\nحساب تو کشتی توست. برای یک سکه، کشتی را غرق نکن.",
        "🏴‍☠️ نقشه گنج\n\nکانال، میانگین، روند، حجم، حد زیان و خروج سریع نقشه گنج معامله‌اند.",
        "🌊 اقیانوس\n\nبازار مثل اقیانوس است. کسی با دست و پا زدن، اقیانوس را شکست نداده.",
        "🏴‍☠️ جزیره گنج\n\nباید با احترام وارد شوی، برداری و برگردی. نباید در جزیره گنج چادر بزنی."
    ],
    "setup": [
        "📊 ستاپ کانال و میانگین\n\n۱. جهت تایم بالا مشخص شود.\n۲. کانال رسم شود.\n۳. میانگین‌ها بررسی شود.\n۴. ورود فقط روی اصلاح.\n۵. حد زیان کوچک.\n۶. خروج سریع.",
        "📐 کانال\n\nکف کانال در روند صعودی، محل جستجوی خرید است. سقف کانال در روند نزولی، محل جستجوی فروش است.",
        "〽️ میانگین‌ها\n\nمیانگین‌ها زبان بازارند. واکنش قیمت را ببین؛ حدس نزن.",
        "📈 خرید حرفه‌ای\n\nدر روند صعودی، روی اصلاح و کف کانال خرید کن؛ نه وسط هیجان شمع بلند.",
        "📉 فروش حرفه‌ای\n\nدر روند نزولی، روی اصلاح و سقف کانال فروش کن؛ نه بعد از ریزش شدید.",
        "⚡ ورود درست\n\nورود خوب باید زود خودش را نشان دهد. اگر حرکت نکرد، شاید جای ورود اشتباه بوده."
    ],
    "risk": [
        "⚖️ حجم\n\nنه در یک نقطه، نه در یک قیمت، نه در یک لحظه، همه حجم را وارد نکن.",
        "🛑 حد زیان\n\nحد زیان یعنی: من اشتباه بودن را می‌پذیرم و حسابم را حفظ می‌کنم.",
        "📉 جبران\n\nبرای جبران یک اشتباه، اشتباه بزرگ‌تر نکن.",
        "⚖️ افزایش حجم\n\nهیچ‌وقت به معامله زیان‌ده برای جبران، حجم اضافه نکن.",
        "📉 زیان کوچک\n\nزیان را کوچک ببند. حساب زنده بماند، فرصت دوباره هست.",
        "📈 سود\n\nسود خوب کردی؟ چارت را ببند. برگرد به زندگی."
    ],
    "samsara": [
        "🔁 سامسارا\n\nضرر → ترس → جبران → حجم بیشتر → ضرر بیشتر. چرخه را بازار نساخته؛ ذهن ساخته است.",
        "🔁 توقف\n\nرهایی از چرخه یعنی بعد از ضرر، توقف؛ نه انتقام.",
        "🔁 ذهن قدیمی\n\nهر بار با همان ذهن قبلی وارد شوی، همان اشتباه قبلی دوباره متولد می‌شود.",
        "🕊 رهایی\n\nدیدن، پذیرفتن، عمل، رها کردن.",
        "🕊 مشاهده\n\nقبل از ورود، سه نفس عمیق. بعد فقط یک سؤال: سیستم چه می‌گوید؟"
    ]
}

for _k, _items in EXTRA_MIND_MESSAGES.items():
    if _k in MIND_MESSAGES:
        for _m in _items:
            if _m not in MIND_MESSAGES[_k]:
                MIND_MESSAGES[_k].append(_m)
    else:
        MIND_MESSAGES[_k] = list(_items)



EXTRA_MIND_MESSAGES_V3 = {
    "mind": [
        "🧠 ذهن معامله‌گر\\n\\nبازار در یک ساعت حساب را نابود نمی‌کند؛ ذهن بی‌قرار این کار را می‌کند.",
        "🧠 سیستم یا احساس؟\\n\\nورود باید از سیستم بیاید، نه از عجله، ترس، انتقام یا فریب یک شمع بلند.",
        "🧠 مشاهده\\n\\nپنج دقیقه فقط نگاه کن. نه صعودی بگو، نه نزولی. فقط ببین.",
        "🧠 ذهن داغ\\n\\nاگر بعد از سود یا ضرر هنوز ذهنت داغ است، معامله ممنوع."
    ],
    "walmart": [
        "🏪 والمارت خانگی\\n\\nتو یک والمارت خانگی داری. والمارت همه موجودی‌اش را در یک ساعت نمی‌فروشد.",
        "🏪 مغازه حساب\\n\\nحساب معاملاتی تو مغازه توست. در یک ساعت مغازه را آتش نزن. فردا هم مشتری هست.",
        "🏪 حفظ مغازه\\n\\nهدف اول فروش بیشتر نیست؛ هدف اول باز ماندن مغازه است."
    ],
    "business": [
        "💼 بیزینس ترید\\n\\nترید یک بیزینس است. هر معامله یک تصمیم مدیریتی است، نه هیجان لحظه‌ای.",
        "💼 گردش بزرگ\\n\\nروزی یک تا ده میلیون دلار می‌توانی معامله کنی؛ اما نه در یک نقطه، نه در یک قیمت، نه در یک لحظه.",
        "💼 اصل بقا\\n\\nاول بقا، بعد سود، بعد رشد حساب."
    ],
    "treasure": [
        "🏴‍☠️ گنج بازار\\n\\nبازار یک گنج عظیم است؛ اما فقط برای کسی که نقشه دارد.",
        "🏴‍☠️ کشتی\\n\\nحساب تو کشتی توست. برای یک سکه، کشتی را غرق نکن.",
        "🏴‍☠️ نقشه گنج\\n\\nکانال، میانگین، روند، حجم، حد زیان و خروج سریع نقشه گنج معامله‌اند."
    ],
    "volume": [
        "⚖️ حجم\\n\\nنه در یک نقطه، نه در یک قیمت، نه در یک لحظه، همه حجم را وارد نکن.",
        "⚖️ تقسیم حجم\\n\\nاگر قرار است سه لات معامله کنی، آن را در چند نقطه منطقی پخش کن.",
        "⚖️ حجم و ذهن\\n\\nحجم سنگین با ذهن آشفته، مثل رانندگی با سرعت بالا و بدون ترمز است."
    ],
    "profit": [
        "📈 سود\\n\\nسود خوب کردی؟ چارت را ببند. برگرد به زندگی.",
        "📈 خروج\\n\\nخروج خوب گاهی مهم‌تر از ورود خوب است."
    ],
    "loss": [
        "📉 زیان\\n\\nزیان امروز را به فاجعه امروز تبدیل نکن. فردا هم هست.",
        "📉 انتقام ممنوع\\n\\nاگر امروز زیان کردی، لازم نیست همین امروز انتقام بگیری."
    ],
    "patience": [
        "⏳ صبر\\n\\nفردا هم بازار هست. هفته بعد هم هست. سال بعد هم هست.",
        "⏳ معامله نکردن\\n\\nمعامله نکردن هم یک معامله است؛ معامله حفظ سرمایه."
    ],
    "checklist": [
        "✅ چک‌لیست ورود\\n\\n۱. جهت تایم بالا مشخص است؟\\n۲. کانال رسم شده؟\\n۳. میانگین‌ها تأیید می‌کنند؟\\n۴. ورود روی اصلاح است؟\\n۵. حد زیان کوچک است؟\\n۶. حجم مناسب است؟",
        "✅ قبل از ورود\\n\\nآیا این معامله حفظ کسب‌وکار است یا تلاش برای جبران؟"
    ]
}

for _k, _items in EXTRA_MIND_MESSAGES_V3.items():
    if _k not in MIND_MESSAGES:
        MIND_MESSAGES[_k] = []
    for _m in _items:
        if _m not in MIND_MESSAGES[_k]:
            MIND_MESSAGES[_k].append(_m)


def make_id(prefix="id"):
    return f"{prefix}_{int(now().timestamp()*1000)}"


def ensure_note_ids():
    notes = load_notes()
    changed = False
    for n in notes:
        if "id" not in n:
            n["id"] = make_id("note")
            changed = True
        if "payload" not in n:
            n["payload"] = {}
            changed = True
        if "time" not in n:
            n["time"] = now_iso()
            changed = True
    if changed:
        save_notes(notes)
    return notes


def find_note(note_id):
    notes = ensure_note_ids()
    for i, n in enumerate(notes):
        if n.get("id") == note_id:
            return notes, i, n
    return notes, None, None


def reminder_summary(rem):
    rt = rem.get("repeat_type", "once")
    rv = rem.get("repeat_value", 1)
    if rt == "once":
        rep = "یک‌بار"
    elif rt == "daily":
        rep = "روزانه"
    elif rt == "weekly":
        rep = "هفتگی"
    elif rt == "monthly":
        rep = "ماهانه"
    elif rt == "every_n_days":
        rep = f"هر {rv} روز"
    elif rt == "every_n_hours":
        rep = f"هر {rv} ساعت"
    elif rt == "every_n_minutes":
        rep = f"هر {rv} دقیقه"
    else:
        rep = rt
    end = rem.get("end_date")
    if end:
        rep += f" تا {end}"
    return rep


def advance_professional_reminder(rem):
    if rem.get("repeat_type", "once") == "once":
        rem["done"] = True
        return

    current = parse_dt(rem["next_time"])
    rt = rem.get("repeat_type", "once")
    rv = int(rem.get("repeat_value", 1) or 1)

    if rt == "daily":
        current += timedelta(days=1)
    elif rt == "weekly":
        current += timedelta(weeks=1)
    elif rt == "monthly":
        y, m, d = current.year, current.month + 1, current.day
        if m > 12:
            m = 1
            y += 1
        d = min(d, monthrange(y, m)[1])
        current = current.replace(year=y, month=m, day=d)
    elif rt == "every_n_days":
        current += timedelta(days=rv)
    elif rt == "every_n_hours":
        current += timedelta(hours=rv)
    elif rt == "every_n_minutes":
        current += timedelta(minutes=rv)
    else:
        rem["done"] = True
        return

    end_date = rem.get("end_date")
    if end_date:
        try:
            end_dt = datetime.strptime(end_date, "%Y-%m-%d").date()
            if current.date() > end_dt:
                rem["done"] = True
                return
        except Exception:
            pass

    rem["next_time"] = current.isoformat(timespec="seconds")


def picker_repeat_text(p):
    rt = p.get("repeat_type", "once")
    rv = p.get("repeat_value", 1)
    if rt == "once":
        return "یک‌بار"
    if rt == "daily":
        return "روزانه"
    if rt == "weekly":
        return "هفتگی"
    if rt == "monthly":
        return "ماهانه"
    if rt == "every_n_days":
        return f"هر {rv} روز"
    if rt == "every_n_hours":
        return f"هر {rv} ساعت"
    if rt == "every_n_minutes":
        return f"هر {rv} دقیقه"
    return rt

def now(): return datetime.now(TZ)
def now_iso(): return now().isoformat(timespec="seconds")
def parse_dt(s): return datetime.fromisoformat(s)

def load_json(path, default):
    if not os.path.exists(path):
        save_json(path, default); return json.loads(json.dumps(default, ensure_ascii=False))
    try:
        with open(path, "r", encoding="utf-8") as f: return json.load(f)
    except Exception:
        return json.loads(json.dumps(default, ensure_ascii=False))

def save_json(path, data):
    with open(path, "w", encoding="utf-8") as f: json.dump(data, f, ensure_ascii=False, indent=2)

def load_config(): return load_json(CONFIG_FILE, DEFAULT_CONFIG)
def save_config(data): save_json(CONFIG_FILE, data)
def load_data(): return load_json(DATA_FILE, {"active": {}, "sessions": []})
def save_data(data): save_json(DATA_FILE, data)
def load_notes(): return load_json(NOTES_FILE, [])
def save_notes(data): save_json(NOTES_FILE, data)
def load_reminders(): return load_json(REMINDERS_FILE, [])
def save_reminders(data): save_json(REMINDERS_FILE, data)
def load_states(): return load_json(USER_STATE_FILE, {})
def save_states(data): save_json(USER_STATE_FILE, data)

def safe_key(text):
    base = text.strip().lower().replace(" ", "_").replace("/", "_").replace(":", "_")
    base = "".join(ch for ch in base if ch.isalnum() or ch in "_-")
    return base[:30] or f"item_{int(now().timestamp())}"

def fmt_duration(seconds):
    minutes = int(float(seconds)//60); h = minutes//60; m = minutes%60
    if h and m: return f"{h} ساعت و {m} دقیقه"
    if h: return f"{h} ساعت"
    return f"{m} دقیقه"

def set_pending(user_id, action, payload):
    states = load_states()
    key = str(user_id)
    existing = states.get(key, {})
    existing["action"] = action
    existing["payload"] = payload
    states[key] = existing
    save_states(states)

def pop_pending(user_id):
    states = load_states()
    key = str(user_id)
    existing = states.get(key)
    if not existing:
        return None
    action = existing.pop("action", None)
    payload = existing.pop("payload", None)
    if existing:
        states[key] = existing
    else:
        states.pop(key, None)
    save_states(states)
    if action:
        return {"action": action, "payload": payload or {}}
    return None

def root_menu():
    return InlineKeyboardMarkup([[InlineKeyboardButton("🏴‍☠️ Trading Mind Master", callback_data="root:mind")], [InlineKeyboardButton("🧠 Personal Life OS", callback_data="root:life")]])

def mind_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🧠 ذهن", callback_data="mind:mind"), InlineKeyboardButton("🏪 والمارت", callback_data="mind:walmart")],
        [InlineKeyboardButton("💼 بیزینس", callback_data="mind:business"), InlineKeyboardButton("🏴‍☠️ گنج بازار", callback_data="mind:treasure")],
        [InlineKeyboardButton("📊 ستاپ", callback_data="mind:setup"), InlineKeyboardButton("🛡 ریسک", callback_data="mind:risk")],
        [InlineKeyboardButton("⚖️ حجم", callback_data="mind:volume"), InlineKeyboardButton("📈 سود", callback_data="mind:profit")],
        [InlineKeyboardButton("📉 زیان", callback_data="mind:loss"), InlineKeyboardButton("⏳ صبر", callback_data="mind:patience")],
        [InlineKeyboardButton("🔁 سامسارا", callback_data="mind:samsara"), InlineKeyboardButton("✅ چک‌لیست", callback_data="mind:checklist")],
        [InlineKeyboardButton("🎲 پیام تصادفی", callback_data="mind:random")],
        [InlineKeyboardButton("🏠 منوی اصلی", callback_data="root")]
    ])


def life_menu():
    cfg = load_config(); rows = []
    for k,v in cfg.items(): rows.append([InlineKeyboardButton(v["title"], callback_data=f"act:{k}")])
    rows += [
        [InlineKeyboardButton("📊 گزارش امروز", callback_data="report:today"), InlineKeyboardButton("📅 بازه خاص", callback_data="range:start")],
        [InlineKeyboardButton("📄 Word امروز", callback_data="export:today:docx"), InlineKeyboardButton("📊 Excel امروز", callback_data="export:today:xlsx")],
        [InlineKeyboardButton("⏰ ریمایندر دکمه‌ای", callback_data="reminder:button")],
        [InlineKeyboardButton("📝 یادداشت‌ها", callback_data="notes:menu"), InlineKeyboardButton("⏹ پایان فعال", callback_data="stop:active")],
        [InlineKeyboardButton("➕ افزودن بخش اصلی", callback_data="add:activity"), InlineKeyboardButton("⚙️ مدیریت بخش‌ها", callback_data="manage:activities")],
        [InlineKeyboardButton("🏠 منوی اصلی", callback_data="root")]
    ]
    return InlineKeyboardMarkup(rows)

def subject_menu(activity):
    cfg = load_config(); rows=[]
    for sk,sv in cfg[activity]["subjects"].items(): rows.append([InlineKeyboardButton(sv["title"], callback_data=f"subj:{activity}:{sk}")])
    if activity == "trading": rows.append([InlineKeyboardButton("➕ افزودن نماد", callback_data="add:trading_symbol")])
    rows += [
        [InlineKeyboardButton("➕ افزودن موضوع", callback_data=f"add:subject:{activity}"), InlineKeyboardButton("⚙️ مدیریت موضوع‌ها", callback_data=f"manage:subjects:{activity}")],
        [InlineKeyboardButton("✏️ ویرایش این بخش", callback_data=f"edit:activity:{activity}"), InlineKeyboardButton("🗑 حذف این بخش", callback_data=f"delete_confirm:activity:{activity}")],
        [InlineKeyboardButton("📝 یادداشت برای این بخش", callback_data=f"note:activity:{activity}")],
        [InlineKeyboardButton("⬅️ برگشت", callback_data="root:life")]
    ]
    return InlineKeyboardMarkup(rows)

def sub_menu(activity, subject):
    cfg = load_config(); rows=[]
    for subk, subv in cfg[activity]["subjects"][subject]["subs"].items(): rows.append([InlineKeyboardButton(subv, callback_data=f"item:{activity}:{subject}:{subk}")])
    rows += [
        [InlineKeyboardButton("➕ افزودن زیرموضوع", callback_data=f"add:sub:{activity}:{subject}"), InlineKeyboardButton("⚙️ مدیریت زیرموضوع‌ها", callback_data=f"manage:subs:{activity}:{subject}")],
        [InlineKeyboardButton("✏️ ویرایش موضوع", callback_data=f"edit:subject:{activity}:{subject}"), InlineKeyboardButton("🗑 حذف موضوع", callback_data=f"delete_confirm:subject:{activity}:{subject}")],
        [InlineKeyboardButton("📝 یادداشت برای این موضوع", callback_data=f"note:subject:{activity}:{subject}")],
        [InlineKeyboardButton("⬅️ برگشت", callback_data=f"act:{activity}")]
    ]
    return InlineKeyboardMarkup(rows)

def item_menu(activity, subject, sub):
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("▶️ شروع", callback_data=f"startsession:{activity}:{subject}:{sub}")],
        [InlineKeyboardButton("📝 متن برای این فعالیت", callback_data=f"note_start:{activity}:{subject}:{sub}")],
        [InlineKeyboardButton("✏️ ویرایش زیرموضوع", callback_data=f"edit:sub:{activity}:{subject}:{sub}"), InlineKeyboardButton("🗑 حذف زیرموضوع", callback_data=f"delete_confirm:sub:{activity}:{subject}:{sub}")],
        [InlineKeyboardButton("📝 یادداشت برای این مورد", callback_data=f"note:item:{activity}:{subject}:{sub}")],
        [InlineKeyboardButton("⬅️ برگشت", callback_data=f"subj:{activity}:{subject}")]
    ])

def manage_activities_menu():
    cfg=load_config(); rows=[]
    for k,v in cfg.items(): rows.append([InlineKeyboardButton("✏️ "+v["title"], callback_data=f"edit:activity:{k}"), InlineKeyboardButton("🗑", callback_data=f"delete_confirm:activity:{k}")])
    rows.append([InlineKeyboardButton("⬅️ برگشت", callback_data="root:life")]); return InlineKeyboardMarkup(rows)

def manage_subjects_menu(activity):
    cfg=load_config(); rows=[]
    for sk,sv in cfg[activity]["subjects"].items(): rows.append([InlineKeyboardButton("✏️ "+sv["title"], callback_data=f"edit:subject:{activity}:{sk}"), InlineKeyboardButton("🗑", callback_data=f"delete_confirm:subject:{activity}:{sk}")])
    rows.append([InlineKeyboardButton("⬅️ برگشت", callback_data=f"act:{activity}")]); return InlineKeyboardMarkup(rows)

def manage_subs_menu(activity, subject):
    cfg=load_config(); rows=[]
    for subk,subv in cfg[activity]["subjects"][subject]["subs"].items(): rows.append([InlineKeyboardButton("✏️ "+subv, callback_data=f"edit:sub:{activity}:{subject}:{subk}"), InlineKeyboardButton("🗑", callback_data=f"delete_confirm:sub:{activity}:{subject}:{subk}")])
    rows.append([InlineKeyboardButton("⬅️ برگشت", callback_data=f"subj:{activity}:{subject}")]); return InlineKeyboardMarkup(rows)

def notes_main_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📒 فهرست یادداشت‌ها", callback_data="notes:list")],
        [InlineKeyboardButton("📅 یادداشت‌های امروز", callback_data="notes:today")],
        [InlineKeyboardButton("📆 یادداشت‌های این هفته", callback_data="notes:week")],
        [InlineKeyboardButton("⬅️ برگشت", callback_data="root:life")]
    ])


def notes_list_menu(notes):
    rows = []
    for n in notes[-15:][::-1]:
        label = note_context_label(n)
        short = (n.get("text","")[:24] + "...") if len(n.get("text","")) > 24 else n.get("text","")
        rows.append([InlineKeyboardButton(f"{n.get('time','')[:10]} | {label} | {short}", callback_data=f"noteview:{n['id']}")])
    rows.append([InlineKeyboardButton("⬅️ برگشت", callback_data="root:life")])
    return InlineKeyboardMarkup(rows)


def note_action_menu(note_id):
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("✏️ ویرایش", callback_data=f"noteedit:{note_id}"), InlineKeyboardButton("🗑 حذف", callback_data=f"notedel:{note_id}")],
        [InlineKeyboardButton("⬅️ فهرست یادداشت‌ها", callback_data="notes:list")]
    ])


def reminders_list_menu(reminders):
    rows = []
    for r in reminders[-15:][::-1]:
        status = "✅" if not r.get("done") and r.get("enabled", True) else "⏸"
        txt = r.get("message","")[:24]
        rows.append([InlineKeyboardButton(f"{status} {r.get('next_time','')[:16]} | {txt}", callback_data=f"remview:{r['id']}")])
    rows.append([InlineKeyboardButton("⬅️ برگشت", callback_data="root:life")])
    return InlineKeyboardMarkup(rows)


def reminder_action_menu(rem_id):
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("⏸/▶️ توقف/فعال", callback_data=f"remtoggle:{rem_id}"), InlineKeyboardButton("🗑 حذف", callback_data=f"remdel:{rem_id}")],
        [InlineKeyboardButton("⬅️ فهرست ریمایندرها", callback_data="reminders:list")]
    ])

def delete_confirm_menu(kind, parts):
    back="root:life"
    if kind=="activity": back=f"act:{parts[0]}"
    elif kind=="subject": back=f"subj:{parts[0]}:{parts[1]}"
    elif kind=="sub": back=f"item:{parts[0]}:{parts[1]}:{parts[2]}"
    return InlineKeyboardMarkup([[InlineKeyboardButton("✅ بله، حذف کن", callback_data="delete:"+kind+":"+":".join(parts))], [InlineKeyboardButton("❌ انصراف", callback_data=back)]])

def clamp_date(y,m,d): return y,m,min(d, monthrange(y,m)[1])
def picker_init(kind):
    t=now(); return {"kind":kind,"year":t.year,"month":t.month,"day":t.day,"hour":t.hour,"minute":0,"text":"","repeat_type":"once","repeat_value":1,"end_date":None}
def picker_text(p,title):
    repeat_text = picker_repeat_text(p)
    end_text = p.get("end_date") or "ندارد"
    return (
        f"{title}\n\n"
        f"📅 تاریخ شروع: {p['year']}-{p['month']:02d}-{p['day']:02d}\n"
        f"⏰ ساعت: {p['hour']:02d}:{p['minute']:02d}\n"
        f"🔁 تکرار: {repeat_text}\n"
        f"📅 پایان تکرار: {end_text}\n\n"
        f"📝 متن: {p.get('text') or 'هنوز نوشته نشده'}"
    )

def reminder_picker_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("سال -", callback_data="rpick:year:-"), InlineKeyboardButton("سال +", callback_data="rpick:year:+")],
        [InlineKeyboardButton("ماه -", callback_data="rpick:month:-"), InlineKeyboardButton("ماه +", callback_data="rpick:month:+")],
        [InlineKeyboardButton("روز -", callback_data="rpick:day:-"), InlineKeyboardButton("روز +", callback_data="rpick:day:+")],
        [InlineKeyboardButton("ساعت -", callback_data="rpick:hour:-"), InlineKeyboardButton("ساعت +", callback_data="rpick:hour:+")],
        [InlineKeyboardButton("دقیقه -", callback_data="rpick:minute:-"), InlineKeyboardButton("دقیقه +", callback_data="rpick:minute:+")],
        [InlineKeyboardButton("یک‌بار", callback_data="rmode:once"), InlineKeyboardButton("روزانه", callback_data="rmode:daily")],
        [InlineKeyboardButton("هفتگی", callback_data="rmode:weekly"), InlineKeyboardButton("ماهانه", callback_data="rmode:monthly")],
        [InlineKeyboardButton("هر N روز", callback_data="rmode:every_n_days"), InlineKeyboardButton("هر N ساعت", callback_data="rmode:every_n_hours")],
        [InlineKeyboardButton("N -", callback_data="rnum:-"), InlineKeyboardButton("N +", callback_data="rnum:+")],
        [InlineKeyboardButton("پایان -۷ روز", callback_data="rend:-"), InlineKeyboardButton("پایان +۷ روز", callback_data="rend:+")],
        [InlineKeyboardButton("حذف تاریخ پایان", callback_data="rend:clear")],
        [InlineKeyboardButton("📝 نوشتن متن", callback_data="rpick:text")],
        [InlineKeyboardButton("✅ ثبت ریمایندر", callback_data="rpick:save")],
        [InlineKeyboardButton("📋 مدیریت ریمایندرها", callback_data="reminders:list")],
        [InlineKeyboardButton("⬅️ برگشت", callback_data="root:life")]
    ])

def get_picker(user_id): return load_states().get(str(user_id),{}).get("picker")
def save_picker(user_id,p):
    states=load_states()
    key=str(user_id)
    existing=states.get(key,{})
    existing["picker"]=p
    states[key]=existing
    save_states(states)
def update_picker(p,field,direction):
    step=1 if direction=="+" else -1
    if field=="year": p["year"]+=step
    elif field=="month":
        p["month"]+=step
        if p["month"]>12: p["month"]=1; p["year"]+=1
        if p["month"]<1: p["month"]=12; p["year"]-=1
    elif field=="day":
        p["day"]+=step; last=monthrange(p["year"],p["month"])[1]
        if p["day"]>last: p["day"]=1
        if p["day"]<1: p["day"]=last
    elif field=="hour": p["hour"]=(p["hour"]+step)%24
    elif field=="minute": p["minute"]=(p["minute"]+5*step)%60
    p["year"],p["month"],p["day"]=clamp_date(p["year"],p["month"],p["day"]); return p

def range_init():
    t=now(); return {"start":{"year":t.year,"month":t.month,"day":t.day},"end":{"year":t.year,"month":t.month,"day":t.day},"focus":"start"}
def range_text(r):
    s,e=r["start"],r["end"]; focus="شروع" if r["focus"]=="start" else "پایان"
    return f"📅 گزارش بازه خاص\n\nشروع: {s['year']}-{s['month']:02d}-{s['day']:02d}\nپایان: {e['year']}-{e['month']:02d}-{e['day']:02d}\n\nدر حال تغییر: {focus}"
def range_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("تغییر شروع", callback_data="range:focus:start"), InlineKeyboardButton("تغییر پایان", callback_data="range:focus:end")],
        [InlineKeyboardButton("سال -", callback_data="range:year:-"), InlineKeyboardButton("سال +", callback_data="range:year:+")],
        [InlineKeyboardButton("ماه -", callback_data="range:month:-"), InlineKeyboardButton("ماه +", callback_data="range:month:+")],
        [InlineKeyboardButton("روز -", callback_data="range:day:-"), InlineKeyboardButton("روز +", callback_data="range:day:+")],
        [InlineKeyboardButton("📊 نمایش متن", callback_data="range:show:text")],
        [InlineKeyboardButton("📄 Word", callback_data="range:export:docx"), InlineKeyboardButton("📊 Excel", callback_data="range:export:xlsx")],
        [InlineKeyboardButton("⬅️ برگشت", callback_data="root:life")]
    ])
def get_range(user_id): return load_states().get(str(user_id),{}).get("range")
def save_range(user_id,r):
    states=load_states(); states[str(user_id)]={"action":"range","payload":{},"range":r}; save_states(states)
def update_range(r,field,direction):
    target=r[r["focus"]]; step=1 if direction=="+" else -1
    if field=="year": target["year"]+=step
    elif field=="month":
        target["month"]+=step
        if target["month"]>12: target["month"]=1; target["year"]+=1
        if target["month"]<1: target["month"]=12; target["year"]-=1
    elif field=="day":
        target["day"]+=step; last=monthrange(target["year"],target["month"])[1]
        if target["day"]>last: target["day"]=1
        if target["day"]<1: target["day"]=last
    target["year"],target["month"],target["day"]=clamp_date(target["year"],target["month"],target["day"]); return r

def range_dates(r):
    s,e=r["start"],r["end"]; sd=datetime(s["year"],s["month"],s["day"]).date(); ed=datetime(e["year"],e["month"],e["day"]).date()
    return (sd,ed) if sd<=ed else (ed,sd)
def get_sessions_between(start_date,end_date):
    result=[]
    for s in load_data().get("sessions",[]):
        try:
            d=datetime.strptime(s.get("date",""),"%Y-%m-%d").date()
            if start_date<=d<=end_date: result.append(s)
        except Exception: pass
    return result
def get_notes_between(start_date,end_date):
    result=[]
    for n in load_notes():
        try:
            d=parse_dt(n["time"]).date()
            if start_date<=d<=end_date: result.append(n)
        except Exception: pass
    return result
def summarize_sessions(sessions):
    totals={}
    for s in sessions:
        key=f"{s['activity_title']} / {s['subject_title']} / {s['sub_title']}"; totals[key]=totals.get(key,0)+float(s["duration_seconds"])
    return totals

def note_context_label(note):
    payload = note.get("payload", {})
    parts = payload.get("parts", [])
    cfg = load_config()

    try:
        if len(parts) >= 3:
            activity, subject, sub = parts[0], parts[1], parts[2]
            return f"{cfg[activity]['title']} / {cfg[activity]['subjects'][subject]['title']} / {cfg[activity]['subjects'][subject]['subs'][sub]}"

        if len(parts) >= 2:
            activity, subject = parts[0], parts[1]
            return f"{cfg[activity]['title']} / {cfg[activity]['subjects'][subject]['title']}"

        if len(parts) >= 1:
            activity = parts[0]
            return f"{cfg[activity]['title']}"
    except Exception:
        pass

    return "یادداشت عمومی"


def note_matches_session(note, session):
    payload = note.get("payload", {})
    parts = payload.get("parts", [])

    if not parts:
        return False

    activity = session.get("activity")
    subject = session.get("subject")
    sub = session.get("sub")

    if len(parts) >= 3:
        return parts[0] == activity and parts[1] == subject and parts[2] == sub

    if len(parts) >= 2:
        return parts[0] == activity and parts[1] == subject

    if len(parts) >= 1:
        return parts[0] == activity

    return False


def group_notes_by_context(notes):
    grouped = {}
    for note in notes:
        label = note_context_label(note)
        grouped.setdefault(label, []).append(note)
    return grouped

def report_text_for_dates(start_date,end_date,title):
    sessions=get_sessions_between(start_date,end_date); notes=get_notes_between(start_date,end_date); totals=summarize_sessions(sessions)
    lines=[title,f"از {start_date} تا {end_date}",f"تهیه گزارش: {now().strftime('%Y-%m-%d %H:%M')}","","خلاصه فعالیت‌ها:"]
    if totals:
        for k,sec in totals.items(): lines.append(f"- {k}: {fmt_duration(sec)}")
    else: lines.append("- فعالیتی ثبت نشده است.")
    lines += ["","یادداشت‌ها:"]
    if notes:
        for n in notes: lines.append(f"- {n['time'][:16]} | {n['text']}")
    else: lines.append("- یادداشتی ثبت نشده است.")
    return "\n".join(lines),sessions,notes,totals

def export_txt_report(start_date, end_date, title):
    ensure_export_dir()
    text, sessions, notes, totals = report_text_for_dates(start_date, end_date, title)
    stamp = now().strftime("%Y%m%d_%H%M")
    path = os.path.join(EXPORT_DIR, f"LifeOS_Report_{stamp}.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path

def export_report(start_date,end_date,title,ext):
    os.makedirs(EXPORT_DIR,exist_ok=True); text,sessions,notes,totals=report_text_for_dates(start_date,end_date,title); stamp=now().strftime("%Y%m%d_%H%M")
    if ext=="docx":
        if Document is None: return export_txt_report(start_date, end_date, title)
        path=os.path.join(EXPORT_DIR,f"LifeOS_Report_{stamp}.docx"); doc=Document(); doc.add_heading("Life OS Report",0)
        for line in text.split("\n"): doc.add_paragraph(line)
        doc.save(path); return path
    if ext=="xlsx":
        if Workbook is None: return export_txt_report(start_date, end_date, title)
        path=os.path.join(EXPORT_DIR,f"LifeOS_Report_{stamp}.xlsx"); wb=Workbook(); ws=wb.active; ws.title="Summary"; ws.append(["Item","Duration"])
        for k,sec in totals.items(): ws.append([k,fmt_duration(sec)])
        ws2=wb.create_sheet("Sessions"); ws2.append(["Date","Activity","Subject","Sub","Start","End","Duration"])
        for s in sessions: ws2.append([s.get("date"),s.get("activity_title"),s.get("subject_title"),s.get("sub_title"),s.get("start_time"),s.get("end_time"),fmt_duration(s.get("duration_seconds",0))])
        ws3=wb.create_sheet("Notes"); ws3.append(["Time","Context","Text"])
        for n in notes: ws3.append([n.get("time"), note_context_label(n), n.get("text")])
        wb.save(path); return path
    raise RuntimeError("فرمت پشتیبانی نمی‌شود.")


async def safe_edit_or_reply(query, text, reply_markup=None):
    try:
        await query.message.edit_text(text, reply_markup=reply_markup)
    except Exception:
        await query.message.reply_text(text, reply_markup=reply_markup)


async def silent_answer(query):
    try:
        await query.answer()
    except Exception:
        pass

async def update_same_message(query, text, reply_markup=None):
    try:
        await query.edit_message_text(text=text, reply_markup=reply_markup)
    except Exception:
        try:
            await query.message.edit_text(text=text, reply_markup=reply_markup)
        except Exception:
            pass


async def edit_picker_message(query, text, reply_markup=None):
    """Edit picker message only. Do not send a new message on +/- clicks."""
    try:
        await query.edit_message_text(text=text, reply_markup=reply_markup)
        return True
    except Exception:
        try:
            await query.message.edit_text(text=text, reply_markup=reply_markup)
            return True
        except Exception:
            return False

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("سیستم اصلی:\n" + VERSION_LABEL, reply_markup=root_menu())

async def text_handler(update:Update, context:ContextTypes.DEFAULT_TYPE):
    user_id=update.effective_user.id; text=update.message.text.strip(); state=pop_pending(user_id)
    if not state:
        saved_picker = get_picker(user_id)
        if saved_picker and saved_picker.get("awaiting_text"):
            state = {"action": "reminder_text", "payload": {"picker": saved_picker}}
        else:
            await update.message.reply_text("برای شروع از /start استفاده کن.")
            return
    action=state["action"]; payload=state["payload"]; cfg=load_config()
    if action == "edit_note_text":
        note_id = payload.get("note_id")
        notes, idx, n = find_note(note_id)
        if n:
            n["text"] = text
            n["edited_at"] = now_iso()
            save_notes(notes)
            await update.message.reply_text("✅ یادداشت ویرایش شد.", reply_markup=note_action_menu(note_id))
        else:
            await update.message.reply_text("یادداشت پیدا نشد.", reply_markup=notes_main_menu())
        return
    if action=="reminder_text":
        p = payload.get("picker") or get_picker(user_id) or picker_init("reminder")
        p["text"] = text
        p["awaiting_text"] = False
        save_picker(user_id,p)
        await update.message.reply_text(picker_text(p,"⏰ ساخت ریمایندر"), reply_markup=reminder_picker_menu())
        return
    if action=="add_activity":
        key=safe_key(text); cfg[key]={"title":text,"subjects":{"default":{"title":"عمومی","subs":{"general":"عمومی"}}}}; save_config(cfg); await update.message.reply_text(f"✅ بخش اصلی اضافه شد: {text}", reply_markup=life_menu()); return
    if action=="add_subject":
        activity=payload["activity"]; key=safe_key(text); cfg[activity]["subjects"][key]={"title":text,"subs":{"general":"عمومی"}}; save_config(cfg); await update.message.reply_text(f"✅ موضوع اضافه شد: {text}", reply_markup=subject_menu(activity)); return
    if action=="add_sub":
        activity,subject=payload["activity"],payload["subject"]; key=safe_key(text); cfg[activity]["subjects"][subject]["subs"][key]=text; save_config(cfg); await update.message.reply_text(f"✅ زیرموضوع اضافه شد: {text}", reply_markup=sub_menu(activity,subject)); return
    if action=="add_trading_symbol":
        key=safe_key(text.upper()); cfg["trading"]["subjects"][key]={"title":text.upper(),"subs":{"analysis":"تحلیل","live":"معامله زنده","journal":"ژورنال معامله","backtest":"بک‌تست"}}; save_config(cfg); await update.message.reply_text(f"✅ نماد اضافه شد: {text.upper()}", reply_markup=subject_menu("trading")); return
    if action=="edit_activity": cfg[payload["activity"]]["title"]=text; save_config(cfg); await update.message.reply_text("✅ نام بخش ویرایش شد.", reply_markup=subject_menu(payload["activity"])); return
    if action=="edit_subject": cfg[payload["activity"]]["subjects"][payload["subject"]]["title"]=text; save_config(cfg); await update.message.reply_text("✅ نام موضوع ویرایش شد.", reply_markup=sub_menu(payload["activity"],payload["subject"])); return
    if action=="edit_sub": cfg[payload["activity"]]["subjects"][payload["subject"]]["subs"][payload["sub"]]=text; save_config(cfg); await update.message.reply_text("✅ نام زیرموضوع ویرایش شد.", reply_markup=item_menu(payload["activity"],payload["subject"],payload["sub"])); return
    if action == "note_start":
        notes = load_notes()
        activity = payload["activity"]
        subject = payload["subject"]
        sub = payload["sub"]
        notes.append({
            "id": make_id("note"),
            "time": now_iso(),
            "type": "note_item",
            "payload": {"parts": [activity, subject, sub]},
            "text": text
        })
        save_notes(notes)

        cfg = load_config()
        label = f"{cfg[activity]['title']} / {cfg[activity]['subjects'][subject]['title']} / {cfg[activity]['subjects'][subject]['subs'][sub]}"
        await update.message.reply_text(
            f"✅ متن برای این فعالیت ثبت شد:\n{label}\n\n{text}",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("▶️ شروع همین فعالیت", callback_data=f"startsession:{activity}:{subject}:{sub}")],
                [InlineKeyboardButton("⬅️ برگشت", callback_data=f"subj:{activity}:{subject}")]
            ])
        )
        return

    if action.startswith("note_"):
        notes=ensure_note_ids(); notes.append({"id":make_id("note"),"time":now_iso(),"type":action,"payload":payload,"text":text}); save_notes(notes); await update.message.reply_text("✅ یادداشت ثبت شد.", reply_markup=life_menu()); return

async def button_handler(update:Update, context:ContextTypes.DEFAULT_TYPE):
    q=update.callback_query; await q.answer(); data=q.data; cfg=load_config()

    if data == "notes:menu":
        await q.message.reply_text("📒 مدیریت یادداشت‌ها:", reply_markup=notes_main_menu())
        return

    if data in ["notes:list", "notes:today", "notes:week"]:
        notes = ensure_note_ids()
        if data == "notes:today":
            today = now().date()
            notes = [n for n in notes if parse_dt(n["time"]).date() == today]
        elif data == "notes:week":
            end = now().date()
            start = end - timedelta(days=6)
            notes = [n for n in notes if start <= parse_dt(n["time"]).date() <= end]
        if not notes:
            await q.message.reply_text("یادداشتی ثبت نشده.", reply_markup=notes_main_menu())
        else:
            await q.message.reply_text("📒 فهرست یادداشت‌ها:", reply_markup=notes_list_menu(notes))
        return

    if data.startswith("noteview:"):
        note_id = data.split(":")[1]
        notes, idx, n = find_note(note_id)
        if not n:
            await q.message.reply_text("یادداشت پیدا نشد.", reply_markup=notes_main_menu())
            return
        await q.message.reply_text(
            f"📒 یادداشت\n\n{note_context_label(n)}\n{n.get('time','')[:16]}\n\n{n.get('text','')}",
            reply_markup=note_action_menu(note_id)
        )
        return

    if data.startswith("noteedit:"):
        note_id = data.split(":")[1]
        set_pending(q.from_user.id, "edit_note_text", {"note_id": note_id})
        await q.message.reply_text("متن جدید یادداشت را بنویس:")
        return

    if data.startswith("notedel:"):
        note_id = data.split(":")[1]
        notes, idx, n = find_note(note_id)
        if idx is not None:
            notes.pop(idx)
            save_notes(notes)
        await q.message.reply_text("🗑 یادداشت حذف شد.", reply_markup=notes_main_menu())
        return

    if data.startswith("rmode:"):
        mode = data.split(":")[1]
        p = get_picker(q.from_user.id) or picker_init("reminder")
        p["repeat_type"] = mode
        save_picker(q.from_user.id,p)
        await edit_picker_message(q, picker_text(p, "⏰ ساخت ریمایندر"), reply_markup=reminder_picker_menu())
        return

    if data.startswith("rnum:"):
        direction = data.split(":")[1]
        p = get_picker(q.from_user.id) or picker_init("reminder")
        n = int(p.get("repeat_value", 1) or 1)
        n += 1 if direction == "+" else -1
        p["repeat_value"] = max(1, n)
        save_picker(q.from_user.id,p)
        await edit_picker_message(q, picker_text(p, "⏰ ساخت ریمایندر"), reply_markup=reminder_picker_menu())
        return

    if data.startswith("rend:"):
        action_end = data.split(":")[1]
        p = get_picker(q.from_user.id) or picker_init("reminder")
        if action_end == "clear":
            p["end_date"] = None
        else:
            base = datetime(p["year"], p["month"], p["day"]).date()
            if p.get("end_date"):
                try:
                    base = datetime.strptime(p["end_date"], "%Y-%m-%d").date()
                except Exception:
                    pass
            base += timedelta(days=7 if action_end == "+" else -7)
            start_day = datetime(p["year"], p["month"], p["day"]).date()
            if base < start_day:
                base = start_day
            p["end_date"] = base.strftime("%Y-%m-%d")
        save_picker(q.from_user.id,p)
        await edit_picker_message(q, picker_text(p, "⏰ ساخت ریمایندر"), reply_markup=reminder_picker_menu())
        return

    if data == "reminders:list":
        reminders = load_reminders()
        if not reminders:
            await q.message.reply_text("ریمایندری ثبت نشده.", reply_markup=life_menu())
        else:
            await q.message.reply_text("📋 مدیریت ریمایندرها:", reply_markup=reminders_list_menu(reminders))
        return

    if data.startswith("remview:"):
        rem_id = data.split(":")[1]
        reminders = load_reminders()
        rem = next((r for r in reminders if r.get("id") == rem_id), None)
        if not rem:
            await q.message.reply_text("ریمایندر پیدا نشد.", reply_markup=life_menu())
            return
        await q.message.reply_text(
            f"⏰ ریمایندر\n\n{rem.get('message','')}\n\nبعدی: {rem.get('next_time','')[:16]}\nتکرار: {reminder_summary(rem)}",
            reply_markup=reminder_action_menu(rem_id)
        )
        return

    if data.startswith("remtoggle:"):
        rem_id = data.split(":")[1]
        reminders = load_reminders()
        for r in reminders:
            if r.get("id") == rem_id:
                r["enabled"] = not r.get("enabled", True)
                break
        save_reminders(reminders)
        await q.message.reply_text("وضعیت ریمایندر تغییر کرد.", reply_markup=reminders_list_menu(reminders))
        return

    if data.startswith("remdel:"):
        rem_id = data.split(":")[1]
        reminders = [r for r in load_reminders() if r.get("id") != rem_id]
        save_reminders(reminders)
        await q.message.reply_text("🗑 ریمایندر حذف شد.", reply_markup=reminders_list_menu(reminders))
        return

    if data=="root": await q.message.reply_text("سیستم اصلی:", reply_markup=root_menu()); return
    if data=="root:mind": await q.message.reply_text("🏴‍☠️ Trading Mind Master:", reply_markup=mind_menu()); return
    if data=="root:life": await q.message.reply_text("🧠 Personal Life OS:", reply_markup=life_menu()); return
    if data.startswith("mind:"):
        key=data.split(":")[1]
        if key=="random":
            allm=[m for arr in MIND_MESSAGES.values() for m in arr]; await q.message.reply_text(random.choice(allm), reply_markup=mind_menu())
        else: await q.message.reply_text(random.choice(MIND_MESSAGES[key]), reply_markup=mind_menu())
        return
    if data=="reminder:button":
        p=picker_init("reminder"); save_picker(q.from_user.id,p); await edit_picker_message(q, picker_text(p, "⏰ ساخت ریمایندر"), reply_markup=reminder_picker_menu()); return
    if data.startswith("rpick:"):
        parts=data.split(":"); p=get_picker(q.from_user.id) or picker_init("reminder")
        if len(parts)==3:
            p=update_picker(p,parts[1],parts[2]); save_picker(q.from_user.id,p); await edit_picker_message(q, picker_text(p, "⏰ ساخت ریمایندر"), reply_markup=reminder_picker_menu()); return
        if parts[1]=="text":
            p["awaiting_text"] = True
            save_picker(q.from_user.id,p)
            set_pending(q.from_user.id,"reminder_text",{"picker":p})
            await q.message.reply_text("متن ریمایندر را بنویس:")
            return
        if parts[1]=="save":
            if not p.get("text"): await update_same_message(q, picker_text(p, "⏰ ساخت ریمایندر") + "\n\n⚠️ اول متن ریمایندر را بنویس.", reply_markup=reminder_picker_menu()); return
            dt=datetime(p["year"],p["month"],p["day"],p["hour"],p["minute"],tzinfo=TZ)
            if dt<=now(): await update_same_message(q, picker_text(p, "⏰ ساخت ریمایندر") + "\n\n⚠️ زمان انتخاب‌شده گذشته است. تاریخ یا ساعت را جلو ببر.", reply_markup=reminder_picker_menu()); return
            p["awaiting_text"] = False
            save_picker(q.from_user.id,p)
            reminders=load_reminders(); reminders.append({"id":make_id("rem"),"chat_id":q.message.chat_id,"type":"advanced","repeat_type":p.get("repeat_type","once"),"repeat_value":int(p.get("repeat_value",1) or 1),"end_date":p.get("end_date"),"enabled":True,"message":p["text"],"start_time":dt.isoformat(timespec="seconds"),"next_time":dt.isoformat(timespec="seconds"),"created_at":now_iso()}); save_reminders(reminders)
            await q.message.reply_text(f"✅ ریمایندر ثبت شد:\n{dt.strftime('%Y-%m-%d %H:%M')}\n{p['text']}", reply_markup=life_menu()); return
    if data=="range:start": r=range_init(); save_range(q.from_user.id,r); await edit_picker_message(q, range_text(r), reply_markup=range_menu()); return
    if data.startswith("range:"):
        parts=data.split(":"); r=get_range(q.from_user.id) or range_init()
        if parts[1]=="focus": r["focus"]=parts[2]; save_range(q.from_user.id,r); await edit_picker_message(q, range_text(r), reply_markup=range_menu()); return
        if parts[1] in ["year","month","day"]: r=update_range(r,parts[1],parts[2]); save_range(q.from_user.id,r); await edit_picker_message(q, range_text(r), reply_markup=range_menu()); return
        if parts[1]=="show": s,e=range_dates(r); text,_,_,_=report_text_for_dates(s,e,"گزارش بازه خاص"); await q.message.reply_text(text[:3900], reply_markup=range_menu()); return
        if parts[1]=="export":
            s,e=range_dates(r)
            ext=parts[2]
            try:
                path=export_report(s,e,"گزارش بازه خاص",ext)
                with open(path, "rb") as f:
                    await q.message.reply_document(document=f, filename=os.path.basename(path), caption="✅ خروجی آماده شد.")
            except Exception as ex:
                await q.message.reply_text(f"خطا در خروجی:\n{ex}")
            return
    if data.startswith("act:"): activity=data.split(":")[1]; await q.message.reply_text(f"{cfg[activity]['title']}\nموضوع را انتخاب کن:", reply_markup=subject_menu(activity)); return
    if data.startswith("subj:"): _,activity,subject=data.split(":"); await q.message.reply_text("زیرموضوع را انتخاب کن:", reply_markup=sub_menu(activity,subject)); return
    if data.startswith("item:"): _,activity,subject,sub=data.split(":"); await q.message.reply_text(f"انتخاب شد:\n{cfg[activity]['subjects'][subject]['subs'][sub]}", reply_markup=item_menu(activity,subject,sub)); return

    if data.startswith("note_start:"):
        _, activity, subject, sub = data.split(":")
        set_pending(q.from_user.id, "note_start", {
            "activity": activity,
            "subject": subject,
            "sub": sub,
            "parts": [activity, subject, sub]
        })
        await q.message.reply_text("متن مربوط به همین فعالیت را بنویس:")
        return

    if data.startswith("startsession:"):
        _,activity,subject,sub=data.split(":"); user_id=str(q.from_user.id); store=load_data()
        if user_id in store["active"]:
            a=store["active"][user_id]; await q.message.reply_text(f"⚠️ یک فعالیت فعال داری:\n{a['activity_title']} / {a['subject_title']} / {a['sub_title']}", reply_markup=life_menu()); return
        start_time=now_iso(); active={"activity":activity,"subject":subject,"sub":sub,"activity_title":cfg[activity]["title"],"subject_title":cfg[activity]["subjects"][subject]["title"],"sub_title":cfg[activity]["subjects"][subject]["subs"][sub],"start_time":start_time}; store["active"][user_id]=active; save_data(store)
        await q.message.reply_text(f"▶️ شروع شد:\n{active['activity_title']} / {active['subject_title']} / {active['sub_title']}\n{parse_dt(start_time).strftime('%Y-%m-%d %H:%M')}", reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("⏹ پایان", callback_data="stop:active")],[InlineKeyboardButton("🏠 Life OS", callback_data="root:life")]])); return
    if data=="stop:active":
        user_id=str(q.from_user.id); store=load_data()
        if user_id not in store["active"]: await q.message.reply_text("فعالیت فعالی وجود ندارد.", reply_markup=life_menu()); return
        active=store["active"].pop(user_id); end_time=now_iso(); st,en=parse_dt(active["start_time"]),parse_dt(end_time); dur=(en-st).total_seconds(); session=dict(active); session.update({"end_time":end_time,"duration_seconds":dur,"date":st.strftime("%Y-%m-%d")}); store["sessions"].append(session); save_data(store)
        await q.message.reply_text(f"⏹ پایان یافت:\n{active['activity_title']} / {active['subject_title']} / {active['sub_title']}\nمدت: {fmt_duration(dur)}", reply_markup=life_menu()); return
    if data.startswith("export:"):
        _,period,ext=data.split(":"); e=now().date(); s=e if period=="today" else e-timedelta(days=6)
        try:
            path=export_report(s,e,"گزارش امروز" if period=="today" else "گزارش ۷ روز اخیر",ext)
            with open(path, "rb") as f:
                await q.message.reply_document(document=f, filename=os.path.basename(path), caption="✅ خروجی آماده شد.")
        except Exception as ex: await q.message.reply_text(f"خطا در خروجی:\n{ex}")
        return
    if data.startswith("add:"):
        parts=data.split(":")
        if parts[1]=="activity": set_pending(q.from_user.id,"add_activity",{}); await q.message.reply_text("نام بخش اصلی جدید را بنویس:")
        elif parts[1]=="subject": set_pending(q.from_user.id,"add_subject",{"activity":parts[2]}); await q.message.reply_text("نام موضوع جدید را بنویس:")
        elif parts[1]=="sub": set_pending(q.from_user.id,"add_sub",{"activity":parts[2],"subject":parts[3]}); await q.message.reply_text("نام زیرموضوع جدید را بنویس:")
        elif parts[1]=="trading_symbol": set_pending(q.from_user.id,"add_trading_symbol",{}); await q.message.reply_text("نماد جدید را بنویس. مثال: GBPJPY")
        return
    if data.startswith("edit:"):
        parts=data.split(":")
        if parts[1]=="activity": set_pending(q.from_user.id,"edit_activity",{"activity":parts[2]}); await q.message.reply_text("نام جدید بخش را بنویس:")
        elif parts[1]=="subject": set_pending(q.from_user.id,"edit_subject",{"activity":parts[2],"subject":parts[3]}); await q.message.reply_text("نام جدید موضوع را بنویس:")
        elif parts[1]=="sub": set_pending(q.from_user.id,"edit_sub",{"activity":parts[2],"subject":parts[3],"sub":parts[4]}); await q.message.reply_text("نام جدید زیرموضوع را بنویس:")
        return
    if data.startswith("delete_confirm:"): parts=data.split(":"); await q.message.reply_text("⚠️ مطمئنی می‌خواهی حذف کنی؟", reply_markup=delete_confirm_menu(parts[1],parts[2:])); return
    if data.startswith("delete:"):
        parts=data.split(":"); kind=parts[1]
        if kind=="activity": cfg.pop(parts[2],None); save_config(cfg); await q.message.reply_text("✅ بخش حذف شد.", reply_markup=life_menu())
        elif kind=="subject": cfg[parts[2]]["subjects"].pop(parts[3],None); save_config(cfg); await q.message.reply_text("✅ موضوع حذف شد.", reply_markup=subject_menu(parts[2]))
        elif kind=="sub": cfg[parts[2]]["subjects"][parts[3]]["subs"].pop(parts[4],None); save_config(cfg); await q.message.reply_text("✅ زیرموضوع حذف شد.", reply_markup=sub_menu(parts[2],parts[3]))
        return
    if data.startswith("manage:"):
        parts=data.split(":")
        if parts[1]=="activities": await q.message.reply_text("⚙️ مدیریت بخش‌ها:", reply_markup=manage_activities_menu())
        elif parts[1]=="subjects": await q.message.reply_text("⚙️ مدیریت موضوع‌ها:", reply_markup=manage_subjects_menu(parts[2]))
        elif parts[1]=="subs": await q.message.reply_text("⚙️ مدیریت زیرموضوع‌ها:", reply_markup=manage_subs_menu(parts[2],parts[3]))
        return
    if data.startswith("note:"): parts=data.split(":"); set_pending(q.from_user.id,"note_"+parts[1],{"parts":parts[2:]}); await q.message.reply_text("یادداشتت را بنویس:"); return
    if data=="notes:list":
        notes=load_notes()[-10:]
        if not notes: await q.message.reply_text("یادداشتی ثبت نشده.", reply_markup=life_menu())
        else: await q.message.reply_text("\n".join(["📝 آخرین یادداشت‌ها:\n"]+[f"• {n['time'][:16]} - {n['text']}" for n in notes]), reply_markup=life_menu())
        return
    if data=="report:today": s=e=now().date(); text,_,_,_=report_text_for_dates(s,e,"گزارش امروز"); await q.message.reply_text(text[:3900], reply_markup=life_menu()); return

def advance_reminder(rem): advance_professional_reminder(rem)
async def reminder_checker(context:ContextTypes.DEFAULT_TYPE):
    reminders=load_reminders(); changed=False
    for rem in reminders:
        if rem.get("done") or not rem.get("enabled", True): continue
        if parse_dt(rem["next_time"])<=now():
            await context.bot.send_message(chat_id=rem["chat_id"], text="⏰ یادآوری:\n\n"+rem["message"]); advance_reminder(rem); changed=True
    if changed: save_reminders([r for r in reminders if not r.get("done")])
async def today_command(update:Update, context:ContextTypes.DEFAULT_TYPE):
    s=e=now().date(); text,_,_,_=report_text_for_dates(s,e,"گزارش امروز"); await update.message.reply_text(text[:3900], reply_markup=life_menu())
async def help_command(update:Update, context:ContextTypes.DEFAULT_TYPE): await update.message.reply_text("/start - منوی اصلی\n/today - گزارش امروز\n/help - راهنما")
async def post_init(app:Application):
    await app.bot.set_my_commands([BotCommand("start","منوی اصلی"), BotCommand("today","گزارش امروز"), BotCommand("help","راهنما")]); app.job_queue.run_repeating(reminder_checker, interval=60, first=10)
def main():
    if not TOKEN: raise ValueError("TELEGRAM_BOT_TOKEN تنظیم نشده است.")
    app=Application.builder().token(TOKEN).post_init(post_init).build(); app.add_handler(CommandHandler("start",start)); app.add_handler(CommandHandler("today",today_command)); app.add_handler(CommandHandler("help",help_command)); app.add_handler(CallbackQueryHandler(button_handler)); app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, text_handler)); print("Life OS with date picker and full mindset is running..."); app.run_polling()
if __name__=="__main__": main()
